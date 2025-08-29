"""
Procesador de Documentos - Servicio Principal COMPLETO
Puerto: 8002
Archivo: services/document-processor/src/main.py
"""

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
import os
import json
import logging
import hashlib
import mimetypes
from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path
import io
import re

# Procesamiento de documentos
try:
    import PyPDF2
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("⚠️ PyPDF2 no disponible - funcionalidad PDF limitada")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("⚠️ python-docx no disponible - funcionalidad DOCX limitada")

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.warning("⚠️ OCR no disponible - instalar tesseract y Pillow")

import pandas as pd
import numpy as np

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ===================================================================
# CONFIGURACIÓN DE DIRECTORIOS
# ===================================================================

# Crear directorios necesarios
BASE_DIR = Path(__file__).parent.parent
UPLOAD_DIR = BASE_DIR / "data" / "uploads"
PROCESSED_DIR = BASE_DIR / "data" / "processed"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
PROCESSED_DIR.mkdir(parents=True, exist_ok=True)

# ===================================================================
# MODELOS DE DATOS
# ===================================================================

class DocumentInfo(BaseModel):
    filename: str
    size: int
    mime_type: str
    document_type: str
    hash_md5: str

class ProcessingOptions(BaseModel):
    extract_text: bool = True
    extract_metadata: bool = True
    extract_tables: bool = False
    perform_ocr: bool = False
    detect_language: bool = True

# ===================================================================
# APLICACIÓN FASTAPI
# ===================================================================

app = FastAPI(
    title="📄 Procesador de Documentos - Agente IA OyP 6.0",
    description="Procesamiento inteligente de documentos PDF, DOCX, TXT, imágenes y más",
    version="6.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Configurar CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===================================================================
# CLASE PRINCIPAL DEL SERVICIO
# ===================================================================

class DocumentProcessorService:
    """Procesador de documentos con capacidades completas"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
            'image/jpeg': 'jpg',
            'image/png': 'png',
            'image/tiff': 'tiff',
            'text/csv': 'csv',
            'application/vnd.ms-excel': 'xls',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx'
        }
        self.processed_documents = {}
        logger.info("✅ Document Processor Service inicializado")
    
    # ===================================================================
    # PROCESAMIENTO PRINCIPAL
    # ===================================================================
    
    async def process_document(self, file: UploadFile, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar documento completo"""
        try:
            # Guardar archivo temporalmente
            file_path = await self.save_uploaded_file(file)
            
            # Obtener información del documento
            doc_info = await self.get_document_info(file_path, file)
            
            # Procesar según el tipo
            doc_type = doc_info['document_type']
            
            if doc_type == 'pdf':
                result = await self.process_pdf(file_path, options)
            elif doc_type == 'docx':
                result = await self.process_docx(file_path, options)
            elif doc_type == 'txt':
                result = await self.process_text_file(file_path, options)
            elif doc_type in ['jpg', 'png', 'tiff']:
                result = await self.process_image(file_path, options)
            elif doc_type in ['csv', 'xls', 'xlsx']:
                result = await self.process_spreadsheet(file_path, options)
            else:
                result = await self.process_unknown_format(file_path, options)
            
            # Combinar información
            result.update(doc_info)
            result['processing_timestamp'] = datetime.now().isoformat()
            
            # Guardar en caché
            self.processed_documents[doc_info['hash_md5']] = result
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando documento: {e}")
            raise HTTPException(status_code=500, detail=f"Error procesando: {str(e)}")
    
    async def save_uploaded_file(self, file: UploadFile) -> Path:
        """Guardar archivo subido"""
        try:
            # Crear nombre único
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename)
            unique_filename = f"{timestamp}_{safe_filename}"
            
            file_path = UPLOAD_DIR / unique_filename
            
            # Guardar archivo
            content = await file.read()
            with open(file_path, "wb") as f:
                f.write(content)
            
            logger.info(f"📁 Archivo guardado: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"❌ Error guardando archivo: {e}")
            raise HTTPException(status_code=500, detail=f"Error guardando: {str(e)}")
    
    async def get_document_info(self, file_path: Path, file: UploadFile) -> Dict[str, Any]:
        """Obtener información del documento"""
        try:
            # Estadísticas del archivo
            file_stat = file_path.stat()
            
            # MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if not mime_type:
                mime_type = file.content_type or 'application/octet-stream'
            
            # Tipo de documento
            document_type = self.supported_formats.get(mime_type, 'unknown')
            
            # Hash MD5
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
            
            return {
                'filename': file.filename,
                'size': file_stat.st_size,
                'mime_type': mime_type,
                'document_type': document_type,
                'hash_md5': file_hash,
                'upload_timestamp': datetime.fromtimestamp(file_stat.st_ctime).isoformat()
            }
            
        except Exception as e:
            logger.error(f"❌ Error obteniendo info: {e}")
            raise HTTPException(status_code=500, detail=f"Error info: {str(e)}")
    
    # ===================================================================
    # PROCESAMIENTO PDF
    # ===================================================================
    
    async def process_pdf(self, file_path: Path, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar archivo PDF"""
        try:
            result = {
                'content_type': 'pdf',
                'text_content': '',
                'metadata': {},
                'pages': 0,
                'tables': [],
                'processing_method': 'PyPDF2' if PDF_AVAILABLE else 'fallback'
            }
            
            if PDF_AVAILABLE:
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    result['pages'] = len(pdf_reader.pages)
                    
                    # Extraer metadatos
                    if options.extract_metadata and pdf_reader.metadata:
                        metadata = pdf_reader.metadata
                        result['metadata'] = {
                            'title': str(metadata.get('/Title', '')),
                            'author': str(metadata.get('/Author', '')),
                            'subject': str(metadata.get('/Subject', '')),
                            'creator': str(metadata.get('/Creator', '')),
                            'producer': str(metadata.get('/Producer', '')),
                            'creation_date': str(metadata.get('/CreationDate', '')),
                            'modification_date': str(metadata.get('/ModDate', ''))
                        }
                    
                    # Extraer texto
                    if options.extract_text:
                        text_parts = []
                        for page_num, page in enumerate(pdf_reader.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text.strip():
                                    text_parts.append(page_text.strip())
                            except Exception as e:
                                logger.warning(f"⚠️ Error página {page_num + 1}: {e}")
                        
                        result['text_content'] = '\n\n'.join(text_parts)
            else:
                result['text_content'] = 'PDF processing no disponible - instalar PyPDF2'
                result['pages'] = 1
            
            # Estadísticas del texto
            if result['text_content']:
                result['statistics'] = self.calculate_text_statistics(result['text_content'])
                
                if options.detect_language:
                    result['language'] = self.detect_language(result['text_content'])
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando PDF: {e}")
            return {
                'content_type': 'pdf',
                'error': str(e),
                'text_content': '',
                'processing_method': 'error'
            }
    
    # ===================================================================
    # PROCESAMIENTO DOCX
    # ===================================================================
    
    async def process_docx(self, file_path: Path, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar archivo DOCX"""
        try:
            result = {
                'content_type': 'docx',
                'text_content': '',
                'metadata': {},
                'paragraphs': [],
                'tables': [],
                'processing_method': 'python-docx' if DOCX_AVAILABLE else 'fallback'
            }
            
            if DOCX_AVAILABLE:
                doc = DocxDocument(file_path)
                
                # Extraer metadatos
                if options.extract_metadata:
                    props = doc.core_properties
                    result['metadata'] = {
                        'title': props.title or '',
                        'author': props.author or '',
                        'subject': props.subject or '',
                        'created': props.created.isoformat() if props.created else '',
                        'modified': props.modified.isoformat() if props.modified else '',
                        'revision': props.revision or 0
                    }
                
                # Extraer texto por párrafos
                if options.extract_text:
                    paragraphs = []
                    for i, paragraph in enumerate(doc.paragraphs):
                        if paragraph.text.strip():
                            paragraphs.append({
                                'index': i,
                                'text': paragraph.text.strip()
                            })
                    
                    result['paragraphs'] = paragraphs
                    result['text_content'] = '\n\n'.join([p['text'] for p in paragraphs])
                
                # Extraer tablas
                if options.extract_tables and doc.tables:
                    tables = []
                    for table_idx, table in enumerate(doc.tables):
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                row_data.append(cell.text.strip())
                            table_data.append(row_data)
                        
                        tables.append({
                            'table_index': table_idx,
                            'rows': len(table_data),
                            'columns': len(table_data[0]) if table_data else 0,
                            'data': table_data
                        })
                    
                    result['tables'] = tables
            else:
                # Fallback: leer como texto plano
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    result['text_content'] = f.read()
                result['processing_method'] = 'text_fallback'
            
            # Estadísticas
            if result['text_content']:
                result['statistics'] = self.calculate_text_statistics(result['text_content'])
                
                if options.detect_language:
                    result['language'] = self.detect_language(result['text_content'])
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando DOCX: {e}")
            return {
                'content_type': 'docx',
                'error': str(e),
                'text_content': '',
                'processing_method': 'error'
            }
    
    # ===================================================================
    # PROCESAMIENTO TEXTO
    # ===================================================================
    
    async def process_text_file(self, file_path: Path, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar archivo de texto"""
        try:
            result = {
                'content_type': 'text',
                'text_content': '',
                'encoding': 'utf-8',
                'lines': [],
                'processing_method': 'text_reader'
            }
            
            # Detectar encoding
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text_content = ''
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    result['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                raise Exception("No se pudo decodificar el archivo")
            
            result['text_content'] = text_content
            
            # Procesar líneas
            if options.extract_text:
                lines = text_content.split('\n')
                result['lines'] = [
                    {'line_number': i + 1, 'content': line}
                    for i, line in enumerate(lines)
                    if line.strip()
                ]
            
            # Estadísticas y detección de idioma
            result['statistics'] = self.calculate_text_statistics(text_content)
            
            if options.detect_language:
                result['language'] = self.detect_language(text_content)
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando texto: {e}")
            return {
                'content_type': 'text',
                'error': str(e),
                'text_content': '',
                'processing_method': 'error'
            }
    
    # ===================================================================
    # PROCESAMIENTO IMÁGENES (OCR)
    # ===================================================================
    
    async def process_image(self, file_path: Path, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar imagen con OCR"""
        try:
            result = {
                'content_type': 'image',
                'text_content': '',
                'image_info': {},
                'ocr_confidence': 0,
                'processing_method': 'tesseract' if OCR_AVAILABLE else 'fallback'
            }
            
            if OCR_AVAILABLE and options.perform_ocr:
                try:
                    # Abrir imagen
                    image = Image.open(file_path)
                    
                    # Información de la imagen
                    result['image_info'] = {
                        'format': image.format,
                        'mode': image.mode,
                        'size': image.size,
                        'width': image.width,
                        'height': image.height
                    }
                    
                    # OCR
                    ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                    
                    # Extraer texto con confianza
                    words = []
                    confidences = []
                    
                    for i in range(len(ocr_data['text'])):
                        if ocr_data['conf'][i] > 30:
                            word = ocr_data['text'][i].strip()
                            if word:
                                words.append(word)
                                confidences.append(ocr_data['conf'][i])
                    
                    result['text_content'] = ' '.join(words)
                    result['ocr_confidence'] = np.mean(confidences) if confidences else 0
                    result['words_detected'] = len(words)
                    
                except Exception as ocr_error:
                    logger.warning(f"⚠️ OCR error: {ocr_error}")
                    result['text_content'] = f"OCR error: {str(ocr_error)}"
                    result['processing_method'] = 'ocr_error'
            else:
                result['text_content'] = 'OCR no disponible - instalar tesseract'
                result['processing_method'] = 'no_ocr'
            
            # Estadísticas si se extrajo texto
            if result['text_content'] and result['text_content'] != 'OCR no disponible - instalar tesseract':
                result['statistics'] = self.calculate_text_statistics(result['text_content'])
                
                if options.detect_language:
                    result['language'] = self.detect_language(result['text_content'])
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando imagen: {e}")
            return {
                'content_type': 'image',
                'error': str(e),
                'text_content': '',
                'processing_method': 'error'
            }
    
    # ===================================================================
    # PROCESAMIENTO HOJAS DE CÁLCULO
    # ===================================================================
    
    async def process_spreadsheet(self, file_path: Path, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar hoja de cálculo"""
        try:
            result = {
                'content_type': 'spreadsheet',
                'text_content': '',
                'sheets': [],
                'tables': [],
                'processing_method': 'pandas'
            }
            
            if file_path.suffix.lower() == '.csv':
                # Procesar CSV
                df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
                sheet_info = {
                    'name': 'Sheet1',
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'data_types': df.dtypes.astype(str).to_dict()
                }
                
                if options.extract_tables:
                    sheet_info['sample_data'] = df.head(10).values.tolist()
                
                result['sheets'] = [sheet_info]
                
            else:
                # Procesar Excel
                try:
                    excel_file = pd.ExcelFile(file_path)
                    
                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        
                        sheet_info = {
                            'name': sheet_name,
                            'rows': len(df),
                            'columns': len(df.columns),
                            'column_names': df.columns.tolist(),
                            'data_types': df.dtypes.astype(str).to_dict()
                        }
                        
                        if options.extract_tables:
                            sheet_info['sample_data'] = df.head(10).values.tolist()
                        
                        result['sheets'].append(sheet_info)
                        
                except Exception as excel_error:
                    logger.warning(f"⚠️ Error Excel: {excel_error}")
                    result['error'] = f"Error procesando Excel: {str(excel_error)}"
            
            # Generar resumen textual
            text_parts = []
            for sheet in result['sheets']:
                text_parts.append(f"Hoja: {sheet['name']}")
                text_parts.append(f"Filas: {sheet['rows']}, Columnas: {sheet['columns']}")
                text_parts.append(f"Columnas: {', '.join(sheet['column_names'][:10])}")
                text_parts.append("")
            
            result['text_content'] = '\n'.join(text_parts)
            
            # Estadísticas
            result['statistics'] = {
                'total_sheets': len(result['sheets']),
                'total_rows': sum(sheet['rows'] for sheet in result['sheets']),
                'total_columns': sum(sheet['columns'] for sheet in result['sheets']),
                'has_data': any(sheet['rows'] > 0 for sheet in result['sheets'])
            }
            
            return result
            
        except Exception as e:
            logger.error(f"❌ Error procesando spreadsheet: {e}")
            return {
                'content_type': 'spreadsheet',
                'error': str(e),
                'text_content': '',
                'processing_method': 'error'
            }
    
    async def process_unknown_format(self, file_path: Path, options: ProcessingOptions) -> Dict[str, Any]:
        """Procesar formato desconocido"""
        try:
            # Intentar leer como texto
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                result = {
                    'content_type': 'unknown_text',
                    'text_content': content,
                    'processing_method': 'text_fallback'
                }
                
                result['statistics'] = self.calculate_text_statistics(content)
                
                if options.detect_language:
                    result['language'] = self.detect_language(content)
                
                return result
                
            except:
                # Leer como binario
                with open(file_path, 'rb') as f:
                    binary_content = f.read()
                
                return {
                    'content_type': 'unknown_binary',
                    'text_content': f'Archivo binario de {len(binary_content)} bytes',
                    'binary_size': len(binary_content),
                    'processing_method': 'binary_inspection'
                }
                
        except Exception as e:
            logger.error(f"❌ Error procesando formato desconocido: {e}")
            return {
                'content_type': 'unknown',
                'error': str(e),
                'text_content': '',
                'processing_method': 'error'
            }
    
    # ===================================================================
    # UTILIDADES
    # ===================================================================
    
    def calculate_text_statistics(self, text: str) -> Dict[str, Any]:
        """Calcular estadísticas del texto"""
        try:
            words = text.split()
            lines = text.split('\n')
            sentences = text.split('.')
            paragraphs = text.split('\n\n')
            
            return {
                'character_count': len(text),
                'character_count_no_spaces': len(text.replace(' ', '')),
                'word_count': len(words),
                'line_count': len(lines),
                'sentence_count': len([s for s in sentences if s.strip()]),
                'paragraph_count': len([p for p in paragraphs if p.strip()]),
                'avg_words_per_sentence': len(words) / max(len([s for s in sentences if s.strip()]), 1),
                'avg_chars_per_word': len(text.replace(' ', '')) / max(len(words), 1),
                'reading_time_minutes': len(words) / 200,  # 200 WPM
                'complexity': 'Alta' if len([w for w in words if len(w) > 6]) / max(len(words), 1) > 0.3 else 'Media'
            }
        except Exception as e:
            logger.error(f"❌ Error calculando estadísticas: {e}")
            return {}
    
    def detect_language(self, text: str) -> Dict[str, Any]:
        """Detección básica de idioma"""
        try:
            spanish_words = ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'le', 'da', 'su']
            english_words = ['the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'i', 'it', 'for', 'not', 'on']
            french_words = ['le', 'de', 'et', 'à', 'un', 'il', 'être', 'en', 'avoir', 'que', 'pour', 'dans', 'ce']
            
            text_lower = text.lower()
            words = text_lower.split()
            
            spanish_count = sum(1 for word in spanish_words if word in words)
            english_count = sum(1 for word in english_words if word in words)
            french_count = sum(1 for word in french_words if word in words)
            
            total_words = len(words)
            
            if spanish_count > english_count and spanish_count > french_count:
                return {'language': 'es', 'confidence': spanish_count / max(total_words, 1)}
            elif english_count > french_count:
                return {'language': 'en', 'confidence': english_count / max(total_words, 1)}
            elif french_count > 0:
                return {'language': 'fr', 'confidence': french_count / max(total_words, 1)}
            else:
                return {'language': 'unknown', 'confidence': 0.0}
                
        except Exception as e:
            logger.error(f"❌ Error detectando idioma: {e}")
            return {'language': 'unknown', 'confidence': 0.0}

# ===================================================================
# INSTANCIA GLOBAL DEL SERVICIO
# ===================================================================

doc_service = DocumentProcessorService()

# ===================================================================
# ENDPOINTS DE LA API
# ===================================================================

@app.get("/")
async def root():
    """Endpoint raíz con información del servicio"""
    return {
        "message": "📄 Bienvenido al Procesador de Documentos - Agente IA OyP 6.0",
        "service": "document-processor",
        "version": "6.0.0",
        "status": "active",
        "supported_formats": list(doc_service.supported_formats.keys()),
        "capabilities": [
            "pdf_processing",
            "docx_processing", 
            "text_processing",
            "image_ocr",
            "spreadsheet_processing",
            "metadata_extraction",
            "language_detection",
            "text_statistics"
        ]
    }

@app.get("/health")
async def health_check():
    """Health check del servicio"""
    return {
        "status": "healthy",
        "service": "document-processor",
        "port": 8002,
        "timestamp": datetime.now().isoformat(),
        "dependencies": {
            "pdf": PDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "ocr": OCR_AVAILABLE
        },
        "processed_docs": len(doc_service.processed_documents)
    }

@app.get("/info")
async def service_info():
    """Información detallada del servicio"""
    return {
        "name": "document-processor",
        "description": "Procesador inteligente de documentos con OCR y análisis de contenido",
        "port": 8002,
        "version": "6.0.0",
        "endpoints": {
            "GET /": "Información del servicio",
            "GET /health": "Health check",
            "GET /info": "Información detallada",
            "POST /upload": "Subir y procesar documento",
            "POST /process_text": "Procesar texto directo",
            "GET /document/{hash}": "Obtener documento procesado",
            "GET /documents": "Listar documentos procesados",
            "GET /formats": "Formatos soportados"
        },
        "supported_formats": doc_service.supported_formats
    }

@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    extract_text: bool = True,
    extract_metadata: bool = True,
    extract_tables: bool = False,
    perform_ocr: bool = False,
    detect_language: bool = True
):
    """Subir y procesar documento"""
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nombre de archivo requerido")
    
    # Validar tamaño (50MB max)
    if hasattr(file, 'size') and file.size and file.size > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Archivo muy grande (máx 50MB)")
    
    # Crear opciones de procesamiento
    options = ProcessingOptions(
        extract_text=extract_text,
        extract_metadata=extract_metadata,
        extract_tables=extract_tables,
        perform_ocr=perform_ocr,
        detect_language=detect_language
    )
    
    result = await doc_service.process_document(file, options)
    return result

@app.post("/process_text")
async def process_text_content(text: str):
    """Procesar contenido de texto directo"""
    try:
        result = {
            "input_text": text,
            "timestamp": datetime.now().isoformat(),
            "statistics": doc_service.calculate_text_statistics(text),
            "language": doc_service.detect_language(text),
            "processing_method": "direct_text"
        }
        
        return result
        
    except Exception as e:
        logger.error(f"❌ Error procesando texto: {e}")
        raise HTTPException(status_code=500, detail=f"Error procesando texto: {str(e)}")

@app.get("/document/{doc_hash}")
async def get_document_info(doc_hash: str):
    """Obtener información de documento procesado"""
    if doc_hash in doc_service.processed_documents:
        return doc_service.processed_documents[doc_hash]
    else:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

@app.get("/documents")
async def list_documents():
    """Listar todos los documentos procesados"""
    return {
        "total_documents": len(doc_service.processed_documents),
        "documents": [
            {
                "hash": doc_hash,
                "filename": doc_info.get("filename", "unknown"),
                "type": doc_info.get("document_type", "unknown"),
                "size": doc_info.get("size", 0),
                "processed": doc_info.get("processing_timestamp", "unknown")
            }
            for doc_hash, doc_info in doc_service.processed_documents.items()
        ]
    }

@app.get("/formats")
async def supported_formats():
    """Listar formatos soportados"""
    return {
        "supported_formats": doc_service.supported_formats,
        "total_formats": len(doc_service.supported_formats),
        "dependencies_status": {
            "pdf_processing": PDF_AVAILABLE,
            "docx_processing": DOCX_AVAILABLE,
            "ocr_processing": OCR_AVAILABLE
        }
    }

@app.delete("/document/{doc_hash}")
async def delete_document(doc_hash: str):
    """Eliminar documento del caché"""
    if doc_hash in doc_service.processed_documents:
        del doc_service.processed_documents[doc_hash]
        return {"message": "Documento eliminado", "hash": doc_hash}
    else:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

@app.get("/statistics")
async def get_processing_statistics():
    """Obtener estadísticas del servicio"""
    docs = doc_service.processed_documents
    
    # Estadísticas por tipo
    type_stats = {}
    for doc in docs.values():
        doc_type = doc.get('document_type', 'unknown')
        type_stats[doc_type] = type_stats.get(doc_type, 0) + 1
    
    # Estadísticas de procesamiento
    processing_stats = {}
    for doc in docs.values():
        method = doc.get('processing_method', 'unknown')
        processing_stats[method] = processing_stats.get(method, 0) + 1
    
    return {
        "total_documents_processed": len(docs),
        "documents_by_type": type_stats,
        "processing_methods": processing_stats,
        "service_uptime": datetime.now().isoformat(),
        "dependencies": {
            "pdf_available": PDF_AVAILABLE,
            "docx_available": DOCX_AVAILABLE,
            "ocr_available": OCR_AVAILABLE
        }
    }

# ===================================================================
# INICIALIZACIÓN DEL SERVICIO
# ===================================================================

if __name__ == "__main__":
    logger.info("🚀 Iniciando Document Processor Service...")
    logger.info("📍 Puerto: 8002")
    logger.info("📖 Documentación: http://localhost:8002/docs")
    logger.info(f"📁 Upload dir: {UPLOAD_DIR}")
    logger.info(f"🔧 PDF: {'✅' if PDF_AVAILABLE else '❌'}")
    logger.info(f"🔧 DOCX: {'✅' if DOCX_AVAILABLE else '❌'}")
    logger.info(f"🔧 OCR: {'✅' if OCR_AVAILABLE else '❌'}")
    
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8002,
        reload=True,
        log_level="info"
    )